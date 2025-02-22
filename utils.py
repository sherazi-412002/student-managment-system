
from docx import Document
import pandas as pd
from io import BytesIO
import json
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.lib.colors import HexColor
from PIL import Image as PILImage

def calculate_grade(percentage):
    """Calculate grade based on percentage"""
    if percentage >= 90:
        return "A+", "Outstanding"
    elif percentage >= 80:
        return "A", "Excellent"
    elif percentage >= 70:
        return "B", "Very Good"
    elif percentage >= 60:
        return "C", "Good"
    elif percentage >= 50:
        return "D", "Satisfactory"
    else:
        return "F", "Fail"

def calculate_percentage(obtained_marks, total_marks):
    """Calculate percentage from obtained and total marks"""
    try:
        return (float(obtained_marks) / float(total_marks)) * 100
    except (ValueError, ZeroDivisionError):
        return 0.0
    

def create_custom_styles():
    """Create custom styles for the PDF"""
    styles = getSampleStyleSheet()
    
    # Title style
    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        textColor=HexColor('#1a237e'),
        alignment=1  # Center alignment
    ))
    
    # Section header style
    styles.add(ParagraphStyle(
        name='SectionHeader',
        parent=styles['Heading2'],
        fontSize=20,
        spaceBefore=15,
        spaceAfter=10,
        textColor=HexColor('#5C6BC0')
    ))
    
    # Info text style
    styles.add(ParagraphStyle(
        name='InfoText',
        parent=styles['Normal'],
        fontSize=14,
        textColor=HexColor('#424242'),
        spaceBefore=6,
        spaceAfter=6
    ))
    
    return styles

def export_to_pdf(student_info, subjects, photo):
    """Export result as PDF using ReportLab with modern design"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=20*mm,
        bottomMargin=20*mm,
        leftMargin=20*mm,
        rightMargin=20*mm
    )
    
    # Get custom styles
    styles = create_custom_styles()
    elements = []
    
    # School/Institution Header
    elements.append(Paragraph("STUDENT RESULT CARD", styles['CustomTitle']))
    elements.append(Spacer(1, 20))
    
    # Student Info Section with Photo
    if photo:
        # Process student photo
        img = PILImage.open(photo)
        img = img.convert('RGB')
        photo_buffer = BytesIO()
        img.save(photo_buffer, format='JPEG')
        photo_buffer.seek(0)
        
        # Create student info layout with photo
        student_photo = Image(photo_buffer, width=1.5*inch, height=1.5*inch)
        
        # Student info in modern card style
        info_elements = [
            [student_photo, 
             Paragraph(f"""
                <font size=16 color="#1a237e"><b>{student_info['name'].upper()}</b></font><br/><br/><br/><br/>
                <font color="#424242">
                Roll Number: {student_info['roll_no']}<br/><br/>
                Class: {student_info['class']}<br/><br/>
                Academic Year: {student_info['academic_year']}
                </font>
             """, styles['InfoText'])]
        ]
        
        info_table = Table(info_elements, colWidths=[1.7*inch, 5*inch])
        info_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('BACKGROUND', (0, 0), (-1, -1), HexColor('#f5f5f5')),
            ('ROUNDEDCORNERS', [10, 10, 10, 10]),
        ]))
        elements.append(info_table)
    
    elements.append(Spacer(1, 20))
    
    # Academic Performance Section
    elements.append(Paragraph("Academic Performance", styles['SectionHeader']))
    
    # Create DataFrame for calculations
    df = pd.DataFrame(subjects)
    df["Percentage"] = df.apply(
        lambda x: calculate_percentage(x["obtained_marks"], x["total_marks"]),
        axis=1
    )
    df["Grade"] = df["Percentage"].apply(lambda x: calculate_grade(x)[0])
    
    # Modern subjects table
    subjects_data = [['Subject', 'Obtained', 'Total', 'Percentage', 'Grade']]
    for _, row in df.iterrows():
        subjects_data.append([
            row['subject'],
            str(row['obtained_marks']),
            str(row['total_marks']),
            f"{row['Percentage']:.1f}%",
            row['Grade']
        ])
    
    subjects_table = Table(subjects_data, colWidths=[2.5*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1*inch])
    subjects_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#E0E0E0')),
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1a237e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('TEXTCOLOR', (0, 1), (-1, -1), HexColor('#424242')),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
        ('TOPPADDING', (0, 1), (-1, -1), 10),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#F5F5F5'), colors.white]),
    ]))
    elements.append(subjects_table)
    elements.append(Spacer(1, 20))
    
    # Overall Performance Section
    total_obtained = df["obtained_marks"].sum()
    total_marks = df["total_marks"].sum()
    overall_percentage = calculate_percentage(total_obtained, total_marks)
    overall_grade, remarks = calculate_grade(overall_percentage)
    
    elements.append(Paragraph("Overall Performance", styles['SectionHeader']))
    
    # Create modern performance cards
    performance_data = [[
        Paragraph(f"""
            <font size=12 color="#1a237e"><b>Total Score</b></font><br/><br/>
            <font size=14 color="#424242">{total_obtained}/{total_marks}</font>
        """, styles['InfoText']),
        Paragraph(f"""
            <font size=12 color="#1a237e"><b>Percentage</b></font><br/><br/>
            <font size=14 color="#424242">{overall_percentage:.1f}%</font>
        """, styles['InfoText']),
        Paragraph(f"""
            <font size=12 color="#1a237e"><b>Grade</b></font><br/><br/>
            <font size=14 color="#424242">{overall_grade}</font>
        """, styles['InfoText'])
    ]]
    
    performance_table = Table(performance_data, colWidths=[2.3*inch, 2.3*inch, 2.3*inch])
    performance_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 25),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 25),
        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#f5f5f5')),
        ('ROUNDEDCORNERS', [10, 10, 10, 10]),
    ]))
    elements.append(performance_table)
    
    # Remarks Section
    elements.append(Spacer(1, 20))
    elements.append(Paragraph(f"""
        <font size=12 color="#1a237e"><b>Remarks:</b></font><br/><br/>
        <font size=11 color="#424242">{remarks}</font>
    """, styles['InfoText']))
    
    # Build PDF
    doc.build(elements)
    pdf_data = buffer.getvalue()
    buffer.close()
    return pdf_data

def export_to_word(student_info, subjects, photo):
    """Export result as Word document"""
    doc = Document()
    
    # Add title
    doc.add_heading('Student Result', 0)
    
    # Add student information
    doc.add_heading('Student Information', level=1)
    doc.add_paragraph(f"Name: {student_info['name']}")
    doc.add_paragraph(f"Roll Number: {student_info['roll_no']}")
    doc.add_paragraph(f"Class: {student_info['class']}")
    doc.add_paragraph(f"Academic Year: {student_info['academic_year']}")
    
    # Add subjects table
    doc.add_heading('Subject Wise Results', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    headers = ['Subject', 'Obtained Marks', 'Total Marks', 'Percentage', 'Grade']
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    # Add subject data
    for subject in subjects:
        percentage = calculate_percentage(subject['obtained_marks'], subject['total_marks'])
        grade = calculate_grade(percentage)[0]
        row_cells = table.add_row().cells
        row_cells[0].text = subject['subject']
        row_cells[1].text = str(subject['obtained_marks'])
        row_cells[2].text = str(subject['total_marks'])
        row_cells[3].text = f"{percentage:.2f}%"
        row_cells[4].text = grade
    
    # Calculate overall results
    total_obtained = sum(subject['obtained_marks'] for subject in subjects)
    total_marks = sum(subject['total_marks'] for subject in subjects)
    overall_percentage = calculate_percentage(total_obtained, total_marks)
    overall_grade, remarks = calculate_grade(overall_percentage)
    
    # Add overall results
    doc.add_heading('Overall Result', level=1)
    doc.add_paragraph(f"Total Marks: {total_obtained}/{total_marks}")
    doc.add_paragraph(f"Overall Percentage: {overall_percentage:.2f}%")
    doc.add_paragraph(f"Overall Grade: {overall_grade}")
    doc.add_paragraph(f"Remarks: {remarks}")
    
    # Save to memory
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

def export_to_json(student_info, subjects):
    """Export result as JSON"""
    # Calculate overall results
    df = pd.DataFrame(subjects)
    total_obtained = df["obtained_marks"].sum()
    total_marks = df["total_marks"].sum()
    overall_percentage = calculate_percentage(total_obtained, total_marks)
    overall_grade, remarks = calculate_grade(overall_percentage)
    
    # Prepare data structure
    result_data = {
        'student_info': student_info,
        'subjects': subjects,
        'overall_result': {
            'total_obtained': int(total_obtained),
            'total_marks': int(total_marks),
            'overall_percentage': float(f"{overall_percentage:.2f}"),
            'overall_grade': overall_grade,
            'remarks': remarks
        }
    }
    
    # Convert to JSON string
    return json.dumps(result_data, indent=2)























